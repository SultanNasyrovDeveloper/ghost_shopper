from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as p_align
from docx.shared import Pt
from num2words import num2words

from django.template.defaultfilters import date as django_date

from ghost_shopper.core.models import CheckKind
from ghost_shopper.check.models import Check

from .base import BaseDocument


MONTH_IN_GENITIVE = {
    1: 'января',
    2: 'февраля',
    3: 'марта',
    4: 'апреля',
    5: 'мая',
    6: 'июня',
    7: 'июля',
    8: 'августа',
    9: 'сентября',
    10: 'октября',
    11: 'ноября',
    12: 'декабря',
}


class AgreementDocument(BaseDocument):
    """
    Organisation agreement document generator.

    Generates agreement document for given organisation, date and checks

    """

    def __init__(self, id, organisation, date, checks):
        """
        Initialize class.
        """
        self.id = id
        self.organisation = organisation
        self.date = date
        self.month_last_day = self._get_months_last_day()
        self.checks = checks
        self.checks_total_price = 0
        self.doc = Document()

    @property
    def file_name(self):
        """
        Generate file name.
        """
        return 'Соглашение_Акт_ТЗ_{}_за_{}.doc'.format(self.organisation.title, self.date.strftime('{}_%Y'.format(
            django_date(self.date, 'F').lower()
        )))

    def generate(self):
        """
        Create organisation documents for given organisation.

        Generates agreement documents for given organisation and saves it.
        """
        self._fill()
        self.doc.save(self.path_to_file)

    def _fill(self):
        """
        Fill agreement document with necessary data.
        """
        self._set_up_doc()
        self._fill_agreement()
        self._fill_task()
        self._fill_act()

    def _set_up_doc(self):
        """
        Set up document.
        """
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

    def _fill_agreement(self):
        """
        Fill documents agreement part.
        """
        self._fill_agreement_header()
        self._fill_agreement_body()
        self._add_signature_table()

    def _fill_agreement_header(self):
        """Fill document agreement part header."""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('Соглашение №{}'.format(self.id).upper())
        run.bold = True
        paragraph.paragraph_format.alignment = p_align.CENTER

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('к Договору {}'.format(self.organisation.contract_name))
        run.italic = True
        paragraph.paragraph_format.alignment = p_align.CENTER

        self._add_empty_line()
        paragraph = self.doc.add_paragraph('«{}» {} {} г.'.format(
            self.month_last_day, MONTH_IN_GENITIVE[self.date.month], self.date.year))
        paragraph.paragraph_format.alignment = p_align.RIGHT
        self._add_empty_line()

        paragraph = self.doc.add_paragraph(
            'Мы, нижеподписавшиеся - "ЗАКАЗЧИК", {}, и "ИСПОЛНИТЕЛЬ", '
            '{} составили настоящее Соглашение, оформляющее выполнение работ в месяце {} {} года.'.format(
                self.organisation.full_name, self.my_organisation.full_name,
                django_date(self.date, 'F').lower(), self.date.strftime('%Y')))
        paragraph.paragraph_format.first_line_indent = Pt(24)
        self._add_empty_line()

    def _fill_agreement_body(self):
        """
        Fill agreement part body.
        """
        paragraph = self.doc.add_paragraph('Исполнителю поручается организация и проведение работ ')
        run = paragraph.add_run('по сбору информации по теме: «{}».'.format(self.organisation.checks_theme))
        run.bold = True
        paragraph.paragraph_format.first_line_indent = Pt(24)

        self.doc.add_paragraph('Работы по настоящему Соглашению выполняются:')
        paragraph = self.doc.add_paragraph('в рамках Договора {} и ТЗ от {} г.'.format(
            self.organisation.contract_name,
            self.date.replace(day=self.month_last_day).strftime('%d.%m.%y'))
        )
        paragraph.paragraph_format.first_line_indent = Pt(24)
        self.doc.add_paragraph('Все работы проводятся на основании Технического Задания к данному Соглашению.')
        self._add_empty_line()

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('1. Сроки выполнения работ:')
        run.bold = True
        paragraph = self.doc.add_paragraph('1.1. Начало выполнения работ: « 01 » {} {} года.'.format(
            MONTH_IN_GENITIVE[self.date.month], self.date.year))
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph = self.doc.add_paragraph('1.1. Окончание выполнения работ: « {} » {} {} года.'.format(
            self.month_last_day, MONTH_IN_GENITIVE[self.date.month], self.date.year))
        paragraph.paragraph_format.first_line_indent = Pt(24)

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('2. Порядок сдачи и приемки работ:')
        run.bold = True
        paragraph = self.doc.add_paragraph(
            '2.1. Приемка и оценка выполненных работ производится по Акту сдачи-приемки работ. ')
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph = self.doc.add_paragraph('2.2. В Акте сдачи-приемки работ отражается объем фактически выполненных '
                                           'Исполнителем работ и производится расчет стоимости выполненных работ.')
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph = self.doc.add_paragraph('2.3. Акт сдачи-приемки работ составляется по окончании выполнения работ. ')
        paragraph.paragraph_format.first_line_indent = Pt(24)

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('3. Стоимость выполняемых работ: ')
        run.bold = True
        paragraph = self.doc.add_paragraph('3.1. Стоимость работ, выполняемых в соответствии с настоящим Соглашением, '
                               'определяется путем умножения удельной стоимости 1 визита на общее количество визитов, '
                               'принятых ЗАКАЗЧИКОМ.')
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph = self.doc.add_paragraph(self._make_check_types_prices_string())
        paragraph.paragraph_format.left_indent = Pt(24)

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('4. Порядок расчетов:')
        run.bold = True
        paragraph = self.doc.add_paragraph('4.1. Аванс не предусмотрен.')
        paragraph.paragraph_format.left_indent = Pt(24)
        paragraph = self.doc.add_paragraph('4.2. Оплата работ производится на основании Акта сдачи-приемки работ.')
        paragraph.paragraph_format.left_indent = Pt(24)
        paragraph = self.doc.add_paragraph('4.3. Оплата производится Заказчиком в течение 15 (Пятнадцати) банковских дней.')
        paragraph.paragraph_format.left_indent = Pt(24)

    def _fill_task_header(self):
        """Fill task part header"""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('ТЕХНИЧЕСКОЕ ЗАДАНИЕ')
        run.bold = True
        paragraph.paragraph_format.alignment = p_align.CENTER

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('на организацию и проведение работ сопровождающих сбор информации')
        run.bold = True
        paragraph.paragraph_format.alignment = p_align.CENTER

        paragraph = self.doc.add_paragraph('К Соглашению №{} от «01» {} {} г.'.format(
            self.id, MONTH_IN_GENITIVE[self.date.month], self.date.year))
        paragraph.paragraph_format.alignment = p_align.CENTER

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('Договор {}'.format(self.organisation.contract_name))
        run.italic = True
        paragraph.paragraph_format.alignment = p_align.CENTER

        self._add_empty_line()
        paragraph = self.doc.add_paragraph('«{}» {} {} г.'.format(
            self.month_last_day, MONTH_IN_GENITIVE[self.date.month], self.date.year))
        paragraph.paragraph_format.alignment = p_align.RIGHT
        self._add_empty_line()

    def _fill_task_body(self):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('1. Содержание работ:')
        run.bold = True

        paragraph = self.doc.add_paragraph(
            '1.1 Исполнитель осуществляет организацию и проведение проверок работы сотрудников в автосалоне Заказчика '
            'в месяце {} {} года методом Mistery Shopping – Тайный Покупатель.'.format(
                django_date(self.date, 'F'), self.date.year))
        paragraph.paragraph_format.left_indent = Pt(24)

        paragraph = self.doc.add_paragraph(
            '1.2. Окончательное количество визитов (проверок) отражается в Акте сдачи-приемки работ как количество '
            'принятых Заказчиком, и служит для расчета стоимости выполненных работ.')
        paragraph.paragraph_format.left_indent = Pt(24)
        self._add_empty_line()

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('2. Порядок проведения работ: ')
        run.bold = True
        self._add_empty_line()

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('2.1. ИСПОЛНИТЕЛЬ осуществляет:')
        run.bold = True
        self._add_empty_line()
        table = self.doc.add_table(rows=3, cols=1)
        table.cell(0, 0).text = '* инструктаж интервьюеров (шопперов);'
        table.cell(0, 1).text = '* сбор информации (методом Mistery Shopping – Тайный Покупатель);'
        table.cell(0, 2).text = '* контроль работы интервьюеров (шопперов) на местах.'
        self._add_empty_line()

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('2.2. ЗАКАЗЧИК предоставляет Исполнителю:')
        run.bold = True
        self._add_empty_line()
        table = self.doc.add_table(rows=2, cols=1)
        table.cell(0, 0).text = '* чек-лист визита;'
        table.cell(1, 0).text = '* выборку с разбивкой по автосалонам'
        self._add_empty_line()

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('3. Результаты работ:')
        run.bold = True
        self._add_empty_line()
        table = self.doc.add_table(rows=2, cols=1)
        table.cell(0, 0).text = '* массив данных в электронном виде (в виде заполненных чек-листов).'
        table.cell(1, 0).text = '* письменный отчёт в электронном виде.'

    def _fill_task(self):
        """ """
        self._fill_task_header()
        self._fill_task_body()
        self._add_signature_table()

    def _fill_act_header(self):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('АКТ СДАЧИ-ПРИЁМКИ РЕЗУЛЬТАТОВ РАБОТ')
        run.bold = True
        paragraph.paragraph_format.alignment = p_align.CENTER

        paragraph = self.doc.add_paragraph('К Соглашению №{} от «01» {} {} г.'.format(
            self.id, MONTH_IN_GENITIVE[self.date.month], self.date.year))
        paragraph.paragraph_format.alignment = p_align.CENTER

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('Договор {}'.format(self.organisation.contract_name))
        run.italic = True
        paragraph.paragraph_format.alignment = p_align.CENTER

        self._add_empty_line()
        paragraph = self.doc.add_paragraph('«{}» {} {} г.'.format(
            self.month_last_day, MONTH_IN_GENITIVE[self.date.month], self.date.year))
        paragraph.paragraph_format.alignment = p_align.RIGHT
        self._add_empty_line()

    def _fill_act_body(self):
        paragraph = self.doc.add_paragraph(
            'Мы, нижеподписавшиеся - "ЗАКАЗЧИК", {}, и "ИСПОЛНИТЕЛЬ", {}, составили настоящий акт о том, что:'.format(
                self.organisation.full_name, self.my_organisation.full_name
            ))
        paragraph.paragraph_format.first_line_indent = Pt(24)
        self._add_empty_line()

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('1. ИСПОЛНИТЕЛЕМ были выполнены работы по сбору '
                               'информации по теме: «{}», а именно:'.format(self.organisation.checks_theme))
        run.bold = True

        paragraph = self.doc.add_paragraph(
            '1.1 Исполнитель осуществляет организацию и проведение проверок работы сотрудников в автосалоне Заказчика '
            'в месяце {} {} года методом Mistery Shopping – Тайный Покупатель.'.format(
                django_date(self.date, 'F'), self.date.year))
        paragraph.paragraph_format.left_indent = Pt(24)

        paragraph = self.doc.add_paragraph(self._make_checks_number_string())
        paragraph.paragraph_format.left_indent = Pt(24)

        paragraph = self.doc.add_paragraph(
            '1.3 Работы выполнены Исполнителем надлежащим образом и в согласованные Сторонами сроки.')
        paragraph.paragraph_format.left_indent = Pt(24)
        self._add_empty_line()

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('2. Стоимость выполненных работ:')
        run.bold = True

        paragraph = self.doc.add_paragraph(
            '2.1 Удельная стоимость 1 посещения дилерского центра, исходя из п.3.2. Соглашения №{} от «01» {} {} года, '
            'составляет:'.format(
                self.id, MONTH_IN_GENITIVE[self.date.month], self.date.year)
        )
        paragraph.paragraph_format.left_indent = Pt(24)
        paragraph = self.doc.add_paragraph(self._make_check_types_prices())
        paragraph.paragraph_format.left_indent = Pt(24)

        paragraph = self.doc.add_paragraph('2.3 Расчет общей стоимости выполненных работ:')
        paragraph.paragraph_format.left_indent = Pt(24)

        paragraph = self.doc.add_paragraph(self._make_check_price_sum_string())
        paragraph.paragraph_format.left_indent = Pt(24)
        self._add_empty_line()

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('3. ЗАКАЗЧИКУ подлежит произвести оплату выполненных Исполнителем работ в размере: '
                               '{} руб. 00 коп. ({} рублей 00 коп.), НДС не облагается.'.format(
            self.checks_total_price, num2words(self.checks_total_price, lang='ru').capitalize()))
        run.bold = True
        self._add_empty_line()

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('4. ЗАКАЗЧИК претензий к качеству выполненных работ не имеет.')
        run.bold = True
        self._add_empty_line()

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('5. Настоящий акт составлен в двух экземплярах, имеющих одинаковую юридическую силу, '
                               'по одному экземпляру для каждой из Сторон.')
        run.bold = True
        self._add_empty_line()


    def _fill_act(self):
        self._fill_act_header()
        self._fill_act_body()
        self._add_signature_table()

    def _add_empty_line(self):
        self.doc.add_paragraph('')

    def _make_check_types_prices_string(self):
        """Make all check types prices string"""
        output = '3.2 '
        for check_kind in CheckKind.objects.all():
            output += 'Удельная стоимость 1 визита с целью "{}" составляет: \n {} руб. 00 коп. ({} рублей 00 коп.)' \
                      'НДС не облагается.\n'.format(
                check_kind.name, check_kind.price, num2words(check_kind.price, lang='ru').capitalize())
        return output

    def _make_check_types_prices(self):
        """"""
        output = ''
        for kind in CheckKind.objects.all():
            output += 'с целью "{}" составляет: \n'.format(kind.name)
            output += '{} руб. 00 коп. ({} рублей 00 коп.) НДС не облагается\n'.format(
                kind.price, num2words(kind.price, lang='ru').capitalize())
        return output

    def _make_check_price_sum_string(self):

        if not self.checks.exists():
            return

        output = ''
        checks_price_sum = 0

        for kind in CheckKind.objects.all():
            kind_checks = self.checks.filter(kind=kind)
            if kind_checks.exists():
                output += ' + {} руб. 00 коп. x {}'.format(kind.price, kind_checks.count())
                checks_price_sum += kind.price * kind_checks.count()

        output += ' = {} руб. 00. ({} рублей 00 коп.), НДС не облагается.'.format(
            checks_price_sum, num2words(checks_price_sum, lang='ru').capitalize())
        self.checks_total_price = checks_price_sum

        return output.strip(' +')

    def _add_signature_table(self):
        self._add_empty_line()
        table = self.doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = '"ЗАКАЗЧИК"'
        table.cell(0, 1).text = '"ИСПОЛНИТЕЛЬ"'
        table.cell(1, 0).text = str(self.organisation.full_name) + ('\n' * 7) + '_' * 15
        table.cell(1, 1).text = self.my_organisation.full_name + ('\n' * 7) + '_' * 15 + '(Юшков Л.В.)'
        self.doc.add_page_break()

    def _make_checks_number_string(self):
        """"""
        output = ''
        output += '1.2 Количество принятых Заказчиком визитов в автосалон - {} ({}), из них: '.format(
            self.checks.count(), num2words(self.checks.count(), lang='ru')
        )
        for kind in CheckKind.objects.all():
            kind_checks = self.checks.filter(kind=kind)
            if kind_checks.count() != 0:
                output += '{} - {}шт., '.format(kind.name, kind_checks.count())
        return output.strip(', ')



